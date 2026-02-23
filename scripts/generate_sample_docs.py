"""Generate sample PDF and DOCX documents for demo purposes.

Usage:
    pip install fpdf2 python-docx
    python scripts/generate_sample_docs.py
"""

import os
import sys

SAMPLES_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "samples")


def generate_hr_policy_pdf(path: str) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, "Employee Handbook 2026", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "1. Annual Leave Policy", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, (
        "All full-time employees are entitled to 15 days of paid annual leave per calendar year. "
        "Leave accrues at 1.25 days per month. Unused leave may be carried over to the next year "
        "up to a maximum of 5 days. Employees must submit leave requests at least 5 business days "
        "in advance through the HR portal. Requests exceeding 5 consecutive days require manager "
        "and HR approval."
    ))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "2. Salary Structure", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, (
        "The company follows a transparent salary banding system. "
        "Below is the salary grid effective January 2026:"
    ))
    pdf.ln(3)

    # Salary table
    headers = ["Grade", "Title", "Base Salary (USD)", "Annual Bonus %", "Stock Options"]
    data = [
        ["L1", "Junior Associate", "55,000 - 65,000", "5%", "None"],
        ["L2", "Associate", "65,000 - 80,000", "8%", "500 units"],
        ["L3", "Senior Associate", "80,000 - 100,000", "12%", "1,000 units"],
        ["L4", "Manager", "100,000 - 130,000", "15%", "2,500 units"],
        ["L5", "Senior Manager", "130,000 - 170,000", "20%", "5,000 units"],
        ["L6", "Director", "170,000 - 220,000", "25%", "10,000 units"],
    ]
    col_widths = [18, 38, 45, 35, 34]

    pdf.set_font("Helvetica", "B", 10)
    for w, h in zip(col_widths, headers):
        pdf.cell(w, 8, h, border=1, align="C")
    pdf.ln()
    pdf.set_font("Helvetica", size=10)
    for row in data:
        for w, val in zip(col_widths, row):
            pdf.cell(w, 7, val, border=1, align="C")
        pdf.ln()
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "3. Remote Work Policy", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, (
        "Employees may work remotely up to 3 days per week with manager approval. "
        "Core collaboration hours are 10:00 AM to 3:00 PM local time. Employees must "
        "maintain a dedicated workspace and stable internet connection (minimum 50 Mbps). "
        "Remote work privileges may be adjusted based on team needs and performance reviews."
    ))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "4. Performance Review Process", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, (
        "Performance reviews are conducted semi-annually (June and December). The process "
        "includes self-assessment, peer feedback (minimum 3 peers), and manager evaluation. "
        "Ratings follow a 5-point scale: Exceptional, Exceeds Expectations, Meets Expectations, "
        "Needs Improvement, Unsatisfactory. Promotions are considered during the December cycle "
        "and require a rating of Exceeds Expectations or above for two consecutive cycles."
    ))

    pdf.output(path)


def generate_legal_policy_docx(path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    title = doc.add_heading("Data Privacy & Compliance Policy", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Effective Date: January 1, 2026 | Version: 3.2 | Classification: Internal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "This policy governs the collection, processing, storage, and disposal of personal data "
        "and confidential business information across all company operations. It applies to all "
        "employees, contractors, vendors, and third-party partners who access company data systems."
    )

    doc.add_heading("2. Data Classification", level=1)
    doc.add_paragraph(
        "All company data must be classified into one of the following categories:"
    )
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Level", "Classification", "Handling Requirements"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows_data = [
        ["1", "Public", "No restrictions; may be shared externally"],
        ["2", "Internal", "Available to all employees; not for external sharing"],
        ["3", "Confidential", "Need-to-know basis; encrypted in transit and at rest"],
        ["4", "Restricted", "Executive approval required; full audit trail mandatory"],
    ]
    for i, row in enumerate(rows_data, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    doc.add_heading("3. Data Retention", level=1)
    doc.add_paragraph(
        "Employee records: 7 years after termination. Financial records: 10 years. "
        "Customer data: duration of contract plus 3 years. Marketing data: 2 years from last "
        "interaction. All data past retention period must be securely deleted within 90 days."
    )

    doc.add_heading("4. Breach Response Protocol", level=1)
    doc.add_paragraph(
        "Any suspected data breach must be reported to the Security Operations Center (SOC) "
        "within 1 hour of discovery. The SOC will initiate the Incident Response Plan, which "
        "includes: (a) containment within 4 hours, (b) impact assessment within 24 hours, "
        "(c) regulatory notification within 72 hours (per GDPR requirements), and "
        "(d) affected party notification within 5 business days."
    )

    doc.add_heading("5. Third-Party Data Sharing", level=1)
    doc.add_paragraph(
        "Sharing data with third parties requires: (1) Data Processing Agreement (DPA) signed "
        "by both parties, (2) Security assessment of the third party completed within 12 months, "
        "(3) Approval from the Data Protection Officer (DPO). No personal data may be transferred "
        "to jurisdictions without adequate data protection unless Standard Contractual Clauses (SCCs) "
        "are in place."
    )

    doc.add_heading("6. Employee Obligations", level=1)
    doc.add_paragraph(
        "All employees must complete annual data privacy training. Employees handling Confidential "
        "or Restricted data must use company-approved devices and VPN. Personal devices may not be "
        "used to store or process data classified Level 3 or above. Violations may result in "
        "disciplinary action up to and including termination."
    )

    doc.save(path)


def generate_finance_report_pdf(path: str) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, "Q4 2025 Financial Report", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Executive Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, (
        "Q4 2025 closed with total revenue of $48.2M, representing 14% year-over-year growth "
        "and exceeding the quarterly target of $45M by 7.1%. Operating expenses were $38.6M, "
        "resulting in an operating margin of 19.9%. Net income reached $7.2M, up from $5.8M in Q4 2024."
    ))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Revenue Breakdown by Division", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    headers = ["Division", "Q4 Revenue", "Q3 Revenue", "QoQ Change", "YoY Change"]
    data = [
        ["Enterprise SaaS", "$22.1M", "$20.5M", "+7.8%", "+18.2%"],
        ["Professional Services", "$12.4M", "$11.8M", "+5.1%", "+11.3%"],
        ["SMB Products", "$8.9M", "$8.2M", "+8.5%", "+9.7%"],
        ["Licensing & Other", "$4.8M", "$4.1M", "+17.1%", "+12.5%"],
        ["TOTAL", "$48.2M", "$44.6M", "+8.1%", "+14.0%"],
    ]
    col_widths = [42, 30, 30, 30, 30]

    pdf.set_font("Helvetica", "B", 10)
    for w, h in zip(col_widths, headers):
        pdf.cell(w, 8, h, border=1, align="C")
    pdf.ln()
    pdf.set_font("Helvetica", size=10)
    for row in data:
        style = "B" if row[0] == "TOTAL" else ""
        pdf.set_font("Helvetica", style, 10)
        for w, val in zip(col_widths, row):
            pdf.cell(w, 7, val, border=1, align="C")
        pdf.ln()
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Operating Expenses", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, (
        "Total operating expenses for Q4 were $38.6M. Key categories: "
        "Personnel costs: $22.3M (57.8%), Technology infrastructure: $6.8M (17.6%), "
        "Sales & marketing: $5.2M (13.5%), General & administrative: $4.3M (11.1%). "
        "Headcount increased to 312 FTEs from 298 in Q3."
    ))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Budget Forecast Q1 2026", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 6, (
        "Projected Q1 2026 revenue: $50.5M (+4.8% QoQ). Planned investments include "
        "$2.1M for AI/ML infrastructure, $1.5M for new hire onboarding (target: 25 new FTEs), "
        "and $0.8M for data center expansion. Expected operating margin: 18.5%."
    ))

    pdf.output(path)


def main() -> None:
    os.makedirs(SAMPLES_DIR, exist_ok=True)

    print("Generating sample documents...")

    hr_path = os.path.join(SAMPLES_DIR, "hr_employee_handbook.pdf")
    generate_hr_policy_pdf(hr_path)
    print(f"  Created: {hr_path}")

    legal_path = os.path.join(SAMPLES_DIR, "legal_data_privacy_policy.docx")
    generate_legal_policy_docx(legal_path)
    print(f"  Created: {legal_path}")

    finance_path = os.path.join(SAMPLES_DIR, "finance_q4_2025_report.pdf")
    generate_finance_report_pdf(finance_path)
    print(f"  Created: {finance_path}")

    print(f"\nDone! 3 sample documents generated in {SAMPLES_DIR}/")
    print("Upload them via the KB Management page to populate the knowledge base.")


if __name__ == "__main__":
    main()
